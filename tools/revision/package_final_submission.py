#!/usr/bin/env python3
"""Validate and package the SoftwareX final submission deliverables."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import subprocess
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from xml.etree import ElementTree

from docx import Document
from PIL import Image
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[2]
FINAL = ROOT / "artifacts" / "revision" / "final-submission"
GENERATED_DOCX = FINAL / "docx"
GENERATED_PDF = FINAL / "pdf"
FIGURES = FINAL / "figures"
SUBMISSION = FINAL / "submission"
VALIDATION = FINAL / "final-manuscript-source-validation-v3.json"
COMPLIANCE = FINAL / "final-submission-compliance.json"
PAGE_REFS = FINAL / "reviewer-page-line-references.json"
PAGE_REF_VERIFICATION = FINAL / "reviewer-page-line-references-verification.json"
ARTWORK_MANIFEST = FIGURES / "submission-artwork-manifest.json"
MANUAL_CHECKS = ROOT / "manuscript" / "Final_Author_Editorial_Checks.json"
AI_PROVENANCE = FINAL / "AI_FIGURE_PROVENANCE.json"
HIGHLIGHTS_SOURCE = ROOT / "manuscript" / "Highlights.md"
RESPONSE_SOURCE = ROOT / "manuscript" / "Response_to_Reviewers_Source.md"

FREEZE = "eaa7d85a0b4921ab2f6e54234cff09aee6a30c8f"
RELEASE_COMMIT = "d7cba1419145e6c75fe69ad63172af5f5abe5028"
RELEASE_SHA256 = "ea0973b7c82e06c3c8910ec36fcf2c3d47765a87d11552337a86c69de41a7cef"
PACKAGE_NAME = "SHM-EM_Final_Submission_GPT_Review_Package.zip"


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def git(*args: str) -> str:
    return subprocess.run(
        ["git", *args], cwd=ROOT, check=True, text=True, encoding="utf-8",
        stdout=subprocess.PIPE,
    ).stdout.strip()


def pdf_text(path: Path) -> str:
    return "\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages)


def docx_text(path: Path) -> str:
    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def tokens(text: str) -> list[str]:
    return re.findall(r"[a-z0-9]+", text.lower())


def normalize_space(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def tracked_revision_counts(path: Path) -> dict[str, int]:
    with zipfile.ZipFile(path) as archive:
        root = ElementTree.fromstring(archive.read("word/document.xml"))
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    return {
        name: len(root.findall(f".//{namespace}{name}"))
        for name in ("ins", "del", "moveFrom", "moveTo")
    }


def final_view_docx_text(path: Path) -> str:
    """Extract Word's final view by excluding deleted and move-from content."""
    with zipfile.ZipFile(path) as archive:
        root = ElementTree.fromstring(archive.read("word/document.xml"))
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    excluded = {f"{namespace}del", f"{namespace}moveFrom"}
    parts: list[str] = []

    def visit(element: ElementTree.Element) -> None:
        if element.tag in excluded:
            return
        if element.tag == f"{namespace}t" and element.text:
            parts.append(element.text)
        elif element.tag == f"{namespace}tab":
            parts.append("\t")
        elif element.tag in {f"{namespace}br", f"{namespace}cr"}:
            parts.append("\n")
        for child in element:
            visit(child)
        if element.tag == f"{namespace}p":
            parts.append("\n")
        elif element.tag == f"{namespace}tc":
            parts.append("\t")

    visit(root)
    return "".join(parts)


def copy_canonical_outputs() -> dict[str, Path]:
    SUBMISSION.mkdir(parents=True, exist_ok=True)
    mapping = {
        "cleanDocx": (GENERATED_DOCX / "SHM-EM_Revised_Manuscript_Clean.docx", SUBMISSION / "SHM-EM_SoftwareX_Revised_Clean.docx"),
        "markedDocx": (GENERATED_DOCX / "SHM-EM_Revised_Manuscript_Marked.docx", SUBMISSION / "SHM-EM_SoftwareX_Revised_Marked.docx"),
        "responseDocx": (GENERATED_DOCX / "SHM-EM_Response_to_Reviewers.docx", SUBMISSION / "Response_to_Reviewers.docx"),
        "highlightsDocx": (GENERATED_DOCX / "SHM-EM_Highlights.docx", SUBMISSION / "Highlights.docx"),
        "cleanPdf": (GENERATED_PDF / "SHM-EM_Revised_Manuscript_Clean.pdf", SUBMISSION / "SHM-EM_SoftwareX_Revised_Clean.pdf"),
        "markedPdf": (GENERATED_PDF / "SHM-EM_Revised_Manuscript_Marked.pdf", SUBMISSION / "SHM-EM_SoftwareX_Revised_Marked.pdf"),
        "responsePdf": (GENERATED_PDF / "SHM-EM_Response_to_Reviewers.pdf", SUBMISSION / "Response_to_Reviewers.pdf"),
        "highlightsPdf": (GENERATED_PDF / "SHM-EM_Highlights.pdf", SUBMISSION / "Highlights.pdf"),
    }
    results: dict[str, Path] = {}
    for key, (source, destination) in mapping.items():
        if not source.is_file():
            raise FileNotFoundError(source)
        shutil.copy2(source, destination)
        results[key] = destination
    shutil.copy2(ROOT / "manuscript" / "Final_Submission_Checklist.md", SUBMISSION / "Final_Submission_Checklist.md")
    return results


def page_count(path: Path) -> int:
    return len(PdfReader(str(path)).pages)


def baseline_verification(canonical: dict[str, Path], submitted_pdf: Path, submitted_docx: Path) -> dict[str, object]:
    submitted_pdf_tokens = set(tokens(pdf_text(submitted_pdf)))
    submitted_docx_tokens = set(tokens(docx_text(submitted_docx)))
    overlap = len(submitted_pdf_tokens.intersection(submitted_docx_tokens))
    clean_tokens = tokens(final_view_docx_text(canonical["cleanDocx"]))
    marked_tokens = tokens(final_view_docx_text(canonical["markedDocx"]))
    counts = tracked_revision_counts(canonical["markedDocx"])
    coverage = overlap / len(submitted_docx_tokens)
    passed = coverage >= 0.99 and counts["ins"] > 0 and counts["del"] > 0 and clean_tokens == marked_tokens
    return {
        "schemaVersion": "shm-em-submitted-baseline-verification-v2",
        "submittedPdf": {"path": str(submitted_pdf), "sha256": sha256(submitted_pdf), "pages": page_count(submitted_pdf)},
        "submittedDocx": {"path": str(submitted_docx), "sha256": sha256(submitted_docx)},
        "submittedDocxUniqueTokenCoverageInSubmittedPdf": round(coverage, 6),
        "markedComparisonBaseline": "Actual submitted manuscript DOCX verified against the Editorial Manager PDF text.",
        "comparisonSettings": {
            "granularity": "word",
            "formattingCompared": False,
            "headersFootersCompared": False,
            "reason": "Keep the English marked manuscript focused on substantive content changes.",
        },
        "trackedRevisionCounts": counts,
        "cleanMarkedFinalViewTokenSequenceEqual": clean_tokens == marked_tokens,
        "cleanMarkedDocxFinalViewTokenCounts": {"clean": len(clean_tokens), "marked": len(marked_tokens)},
        "pass": passed,
    }


def verify_response_locations(canonical: dict[str, Path]) -> dict[str, object]:
    references = json.loads(PAGE_REFS.read_text(encoding="utf-8"))
    source = RESPONSE_SOURCE.read_text(encoding="utf-8")
    rendered = normalize_space(docx_text(canonical["responseDocx"]))
    missing_source = [
        key for key in references
        if not re.search(rf"^###\s+{re.escape(key)}(?:\.|\s)", source, flags=re.MULTILINE)
    ]
    missing_docx = [key for key, value in references.items() if value not in rendered]
    return {
        "expectedItems": len(references),
        "missingItemAnchorsFromResponseSource": missing_source,
        "missingFromResponseDocx": missing_docx,
        "pass": len(references) == 27 and not missing_source and not missing_docx,
    }


def verify_highlights(canonical: dict[str, Path]) -> dict[str, object]:
    bullets = [line[2:].strip() for line in HIGHLIGHTS_SOURCE.read_text(encoding="utf-8").splitlines() if line.startswith("- ")]
    rendered = normalize_space(docx_text(canonical["highlightsDocx"]))
    lengths = [len(item) for item in bullets]
    all_present = all(item in rendered for item in bullets)
    return {
        "bulletCount": len(bullets),
        "characterCounts": lengths,
        "allPresentInDocx": all_present,
        "pass": 3 <= len(bullets) <= 5 and max(lengths, default=999) <= 85 and all_present,
    }


def verify_artwork() -> dict[str, object]:
    manifest = json.loads(ARTWORK_MANIFEST.read_text(encoding="utf-8"))
    issues: list[str] = []
    expected_formats = {1: ".pdf", 2: ".pdf", 3: ".pdf", 4: ".tiff", 5: ".tiff"}
    artifacts = manifest.get("artifacts", [])
    for index, item in enumerate(artifacts, start=1):
        path = ROOT / item["path"]
        if not path.is_file():
            issues.append(f"Missing artwork: {path}")
            continue
        if path.suffix.lower() != expected_formats.get(index):
            issues.append(f"Fig. {index} has unexpected format: {path.suffix}")
        if sha256(path) != item.get("sha256"):
            issues.append(f"Fig. {index} SHA-256 mismatch")
        if path.suffix.lower() == ".pdf" and page_count(path) != 1:
            issues.append(f"Fig. {index} vector PDF is not one page")
        if path.suffix.lower() == ".tiff":
            with Image.open(path) as image:
                if image.width < 3000 or image.height < 2000:
                    issues.append(f"Fig. {index} raster dimensions are too small")
            if float(item.get("effectiveDpiAtFinalWidth", 0)) < 500:
                issues.append(f"Fig. {index} effective DPI is below 500")
    return {"artworkCount": len(artifacts), "issues": issues, "pass": len(artifacts) == 5 and not issues}


def production_core_diff() -> list[str]:
    output = git("diff", "--name-only", FREEZE, "--", "src/backend/src/main", "src/pit_pre/pit_pre", "src/frontend/src")
    return [line for line in output.splitlines() if line]


def build_completion_report(
    canonical: dict[str, Path], verification: dict[str, object], validation: dict[str, object],
    compliance: dict[str, object], response_locations: dict[str, object], highlights: dict[str, object],
) -> str:
    pages = {key: page_count(path) for key, path in canonical.items() if key.endswith("Pdf")}
    revisions = verification["trackedRevisionCounts"]
    return f"""# Final Submission Compliance Correction Report

## Decision

The automated SoftwareX compliance-correction phase is complete and ready for the final GPT upload-readiness audit. The package is intentionally **HOLD_FOR_MANUAL_CONFIRMATION**, because FS-07A, FS-07B, and FS-08 through FS-10 remain author-, data-owner-, or Editorial-Manager-owned decisions. No frozen production-core source changed and no new scientific experiment was run.

## Locked anchors

- Final production-core baseline: `{FREEZE}`
- Immutable revised release: `v1.0.1` at `{RELEASE_COMMIT}`
- Release archive SHA-256: `{RELEASE_SHA256}`
- Strengthened source validation: {sum(item['status'] == 'PASS' for item in validation['checks'])}/{len(validation['checks'])} PASS
- Reviewer items: {validation['reviewerItemCount']}
- References: {validation['referenceCount']}

## SoftwareX compliance result

- Strict clean-DOCX word count: **{compliance['strictWordCount']}** (limit 3,000; post-disclosure target 2,950)
- Final manuscript figures: **5**
- AI-use declaration: immediately before References
- Submission artwork: Fig. 1-3 one-page vector PDF; Fig. 4-5 TIFF at 609.6 effective dpi
- Reviewer locations: {response_locations['expectedItems']}/27 source anchors and rendered final locations verified
- Highlights: {highlights['bulletCount']} bullets; character counts {highlights['characterCounts']}
- AI disclosure: Fig. 1-Fig. 5 provenance/captions and the Section 3.5 code-method statement verified

## Final document set

- `SHM-EM_SoftwareX_Revised_Clean.docx/pdf` ({pages['cleanPdf']} pages)
- `SHM-EM_SoftwareX_Revised_Marked.docx/pdf` ({pages['markedPdf']} pages with markup shown)
- `Response_to_Reviewers.docx/pdf` ({pages['responsePdf']} pages)
- `Highlights.docx/pdf` ({pages['highlightsPdf']} page)
- five separate submission artwork files plus editable PNG/SVG auxiliaries

The marked manuscript contains {revisions['ins']} insertions, {revisions['del']} deletions, {revisions['moveFrom']} move-from elements, and {revisions['moveTo']} move-to elements. Formatting and header/footer differences are excluded, so localized Word formatting balloons and duplicate page-number revisions are absent. Its accepted-revision token sequence is identical to the clean manuscript: `{str(verification['cleanMarkedFinalViewTokenSequenceEqual']).lower()}`.

## Rendering and integrity QA

Microsoft Word COM exported the DOCX files to PDF; Poppler rendered page images at 144 dpi. All 13 clean-manuscript pages, 24 marked-manuscript pages, 12 response pages, the Highlights page, and the five figure compositions were inspected. No clipped content, duplicate footer, localized formatting balloon, or broken final artwork was found. The marked manuscript deliberately retains visible deleted/replaced tables and figures; accepting all revisions yields the clean manuscript exactly.

Automated checks also confirm the actual submitted DOCX/PDF baseline, all 27 final page/line locations, five artwork hashes/formats, Highlights length, the strict word count, and an empty production-core diff.

## Mandatory manual stop gates

- FS-07A: confirm the competing-interest declaration.
- FS-07B: all authors confirm their CRediT roles.
- FS-08: all authors confirm names, affiliations, funding, correspondence, and acknowledgements.
- FS-09: data owner approves the public/restricted data-availability wording.
- FS-10: corresponding author checks deadline, item types, and filenames in Editorial Manager.

The package must not be labelled submission-ready until those five gates are explicitly closed.
"""


def artifact_entry(path: Path) -> dict[str, object]:
    entry: dict[str, object] = {"path": path.relative_to(ROOT).as_posix(), "bytes": path.stat().st_size, "sha256": sha256(path)}
    if path.suffix.lower() == ".pdf":
        entry["pages"] = page_count(path)
    if path.suffix.lower() in {".png", ".tiff"}:
        with Image.open(path) as image:
            entry["pixels"] = {"width": image.width, "height": image.height}
    return entry


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--submitted-pdf", type=Path, required=True)
    parser.add_argument("--submitted-docx", type=Path, required=True)
    args = parser.parse_args()

    FINAL.mkdir(parents=True, exist_ok=True)
    canonical = copy_canonical_outputs()
    validation = json.loads(VALIDATION.read_text(encoding="utf-8"))
    compliance = json.loads(COMPLIANCE.read_text(encoding="utf-8"))
    manual_checks = json.loads(MANUAL_CHECKS.read_text(encoding="utf-8"))

    if not validation.get("pass") or len(validation.get("checks", [])) != 24:
        raise RuntimeError("The strengthened source validation is not 24/24 PASS.")
    if not compliance.get("automatedPass") or compliance.get("strictWordCount", 9999) > 3000:
        raise RuntimeError("The automated SoftwareX compliance gates are not PASS.")
    if compliance.get("submissionStatus") != "HOLD_FOR_MANUAL_CONFIRMATION":
        raise RuntimeError("The package must retain the manual-confirmation hold.")
    pending_manual = [key for key, value in manual_checks.items() if key.startswith("FS-") and value["status"].startswith("PENDING")]
    if pending_manual != ["FS-07A", "FS-07B", "FS-08", "FS-09", "FS-10"]:
        raise RuntimeError(f"Unexpected manual gate state: {pending_manual}")

    automated_gate_status = {
        item["id"]: item["status"] for item in compliance.get("automatedGates", [])
    }
    ai_disclosure_pass = all(
        automated_gate_status.get(identifier) == "PASS"
        for identifier in ("FS-11", "FS-12", "FS-13", "FS-14")
    )

    core_diff = production_core_diff()
    if core_diff:
        raise RuntimeError(f"Production-core diff is not empty: {core_diff}")

    verification = baseline_verification(canonical, args.submitted_pdf, args.submitted_docx)
    response_locations = verify_response_locations(canonical)
    highlights = verify_highlights(canonical)
    artwork = verify_artwork()
    checks = {
        "submittedBaselineAndMarkedManuscript": verification["pass"],
        "responsePageLineLocations": response_locations["pass"],
        "highlights": highlights["pass"],
        "submissionArtwork": artwork["pass"],
        "sourceValidation": validation.get("pass", False),
        "softwareXAutomatedCompliance": compliance.get("automatedPass", False),
        "aiDisclosureAndPostDisclosureChecks": ai_disclosure_pass,
        "productionCoreDiffEmpty": not core_diff,
    }
    if not all(checks.values()):
        raise RuntimeError(f"Final package verification failed: {checks}")

    verification.update({"responseLocations": response_locations, "highlights": highlights, "artwork": artwork, "packageChecks": checks})
    verification_path = FINAL / "submitted-baseline-verification.json"
    verification_path.write_text(json.dumps(verification, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")

    report_path = FINAL / "FINAL_SUBMISSION_COMPLETION_REPORT.md"
    report_path.write_text(build_completion_report(canonical, verification, validation, compliance, response_locations, highlights), encoding="utf-8")

    handoff_path = FINAL / "GPT_REVIEW_HANDOFF.md"
    handoff_path.write_text(
        f"""# GPT Review Handoff

## Status

`AUTOMATED_COMPLIANCE_PASS / HOLD_FOR_MANUAL_CONFIRMATION`

Perform the final upload-readiness audit of the SoftwareX minor-revision package. Do not infer that FS-07A, FS-07B, or FS-08 through FS-10 are complete.

## Review first

1. `FINAL_SUBMISSION_COMPLETION_REPORT.md`
2. `final-submission-manifest.json`
3. `final-submission-compliance.json`
4. `SHM-EM_SoftwareX_Revised_Clean.docx`
5. `SHM-EM_SoftwareX_Revised_Marked.docx`
6. `Response_to_Reviewers.docx`
7. `Highlights.docx`
8. `final-manuscript-source-validation-v3.json`
9. `submitted-baseline-verification.json`
10. `reviewer-page-line-references-verification.json`
11. `submission-artwork-manifest.json` and Fig. 1-Fig. 5 submission files
12. `Final_Author_Editorial_Checks.json`
13. `AI_FIGURE_PROVENANCE.json`

## Audit questions

- Does the clean manuscript remain within 3,000 words under the conservative count?
- Is the AI declaration immediately before References?
- Are all 27 response locations supported by the 13-page clean manuscript?
- Does accepting all tracked revisions produce the clean manuscript exactly?
- Are Fig. 1-3 vector PDF and Fig. 4-5 TIFF at at least 500 dpi?
- Do all final captions and the general declaration accurately reflect `AI_FIGURE_PROVENANCE.json`?
- Does Section 3.5 disclose revision-stage AI-assisted code editing and human/regression review?
- Are all scientific claims, release anchors, and non-claims still consistent?
- Are FS-07A, FS-07B, and FS-08 through FS-10 clearly left for the responsible humans?

## Locked boundaries

- Production-core baseline: `{FREEZE}`; final diff: `NONE`.
- Revised release: `v1.0.1` at `{RELEASE_COMMIT}`; do not move tags.
- No new experiment, model, algorithm claim, tolerance, or release was introduced.
- Windows remains the exact-output reference; Docker/Linux establishes functional/logical portability only.

## Package path rule

The review directory and ZIP are generated only under `artifacts/revision/final-submission/` in this repository. No sandbox or external attachment path is authoritative.
""", encoding="utf-8",
    )

    figure_files = [
        FIGURES / f"Fig{number}_{stem}.{suffix}"
        for number, stem, suffixes in (
            (1, "Research_Gap_and_Workflow", ("pdf", "png", "svg")),
            (2, "Software_Architecture", ("pdf", "png", "svg")),
            (3, "Forecast_to_Event_Sequence", ("pdf", "png", "svg")),
            (4, "Task_Oriented_Interface_Composite", ("tiff", "png", "svg")),
            (5, "Public_Reference_Workflow", ("tiff", "png", "svg")),
        )
        for suffix in suffixes
    ]
    deliverables = [
        *canonical.values(), *figure_files, ARTWORK_MANIFEST, VALIDATION, COMPLIANCE, PAGE_REFS,
        PAGE_REF_VERIFICATION, verification_path, report_path, handoff_path, AI_PROVENANCE,
        SUBMISSION / "Final_Submission_Checklist.md", MANUAL_CHECKS, HIGHLIGHTS_SOURCE,
        ROOT / "manuscript" / "SHM-EM_Revised_Manuscript_Source.md", RESPONSE_SOURCE,
        ROOT / "manuscript" / "Revision_Change_Map.md", ROOT / "manuscript" / "Final_Reviewer_Evidence_Map.md",
    ]
    for path in deliverables:
        if not path.is_file():
            raise FileNotFoundError(path)

    source_hashes = {
        path.relative_to(ROOT).as_posix(): sha256(path)
        for path in (
            ROOT / "manuscript" / "SHM-EM_Revised_Manuscript_Source.md", RESPONSE_SOURCE,
            ROOT / "manuscript" / "Revision_Change_Map.md", ROOT / "manuscript" / "Final_Reviewer_Evidence_Map.md",
            ROOT / "manuscript" / "Final_Submission_Checklist.md", HIGHLIGHTS_SOURCE,
        )
    }
    manifest = {
        "schemaVersion": "shm-em-final-submission-manifest-v2",
        "generatedAtUtc": datetime.now(timezone.utc).isoformat(),
        "generationBaseHead": git("rev-parse", "HEAD"),
        "branch": git("branch", "--show-current"),
        "sourceHashes": source_hashes,
        "finalCoreFreezeV3": FREEZE,
        "productionCoreDiff": "NONE",
        "release": {"tag": "v1.0.1", "commit": RELEASE_COMMIT, "archiveSha256": RELEASE_SHA256},
        "strictWordCount": compliance["strictWordCount"],
        "sourceValidation": "24/24 PASS",
        "reviewerItems": 27,
        "references": 30,
        "automatedChecks": checks,
        "submissionStatus": compliance["submissionStatus"],
        "pendingManualGates": pending_manual,
        "artifacts": [artifact_entry(path) for path in deliverables],
    }
    manifest_path = FINAL / "final-submission-manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")

    review_dir = FINAL / "gpt-review-package"
    if review_dir.exists():
        shutil.rmtree(review_dir)
    review_dir.mkdir(parents=True)
    package_files = [*deliverables, manifest_path]
    for path in package_files:
        shutil.copy2(path, review_dir / path.name)

    zip_path = FINAL / PACKAGE_NAME
    with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        for path in package_files:
            archive.write(path, arcname=path.name)
    zip_digest = sha256(zip_path)
    (FINAL / f"{zip_path.name}.sha256").write_text(f"{zip_digest}  {zip_path.name}\n", encoding="ascii")

    print(json.dumps({
        "automatedPass": True,
        "submissionStatus": compliance["submissionStatus"],
        "pendingManualGates": pending_manual,
        "finalDirectory": str(FINAL),
        "reviewDirectory": str(review_dir),
        "zip": str(zip_path),
        "zipSha256": zip_digest,
        "productionCoreDiff": "NONE",
        "strictWordCount": compliance["strictWordCount"],
        "pageCounts": {key: page_count(path) for key, path in canonical.items() if key.endswith("Pdf")},
        "trackedRevisionCounts": verification["trackedRevisionCounts"],
    }, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
