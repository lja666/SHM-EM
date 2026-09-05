#!/usr/bin/env python3
"""Generate the clean manuscript and reviewer-response DOCX artifacts."""

from __future__ import annotations

import argparse
import json
import os
import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DEFAULT_TEMPLATE = (
    Path(os.environ["SHM_EM_SUBMITTED_TEMPLATE"])
    if os.environ.get("SHM_EM_SUBMITTED_TEMPLATE")
    else None
)
FINAL_FIGURES = ROOT / "artifacts" / "revision" / "final-submission" / "figures"

BLUE = "0B3C82"
MID_BLUE = "1769E8"
TEAL = "087F83"
LIGHT_BLUE = "EAF2FB"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "5E6B80"
RED = "B3261E"
BLACK = "172033"


FIGURES = {
    1: (
        FINAL_FIGURES / "Fig1_Research_Gap_and_Workflow.png",
        "Fig. 1. Research gaps, the SHM-EM software boundary, and the forecast-aware user workflow.",
        Mm(165),
    ),
    2: (
        FINAL_FIGURES / "Fig2_Software_Architecture.png",
        "Fig. 2. Four-layer SHM-EM architecture. MySQL is the validated reference persistence implementation; the observation registry and service interfaces define the storage-adapter extension boundary.",
        Mm(165),
    ),
    3: (
        FINAL_FIGURES / "Fig3_Forecast_to_Event_Sequence.png",
        "Fig. 3. Controlled sequence from persisted forecasts through optional Project Future State inspection, audited Evaluate, independently gated Execute, and formal provenance.",
        Mm(175),
    ),
    4: (
        FINAL_FIGURES / "Fig4_Task_Oriented_Interface_Composite.png",
        "Fig. 4. Task-oriented interface views of SHM-EM: (a) project-level observed and forecast risk, (b) a joint engineering-valued observation/forecast series, and (c) prediction-batch completeness and execution eligibility. The interface is illustrative; quantitative validation is reported in the contract, failure-path, runtime, reuse, and provenance evidence.",
        Mm(175),
    ),
    5: (
        FINAL_FIGURES / "Fig5_Public_Reference_Workflow.png",
        "Fig. 5. Public reference case, verified six-model contract, common temporal frame, and end-to-end reproduction checks.",
        Mm(175),
    ),
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value: bool = True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    props.append(size)
    run.append(props)
    fld.append(run)
    paragraph._p.append(fld)


def reset_footer(footer, *, include_page_number: bool) -> None:
    element = footer._element
    for child in list(element):
        element.remove(child)
    paragraph = footer.add_paragraph()
    if include_page_number:
        add_page_number(paragraph)


def configure_document(document: Document, title: str, subject: str) -> None:
    props = document.core_properties
    props.title = title
    props.subject = subject
    props.author = "Ji'an Liao; Zifa Wang; Dengke Zhao; Jianming Wang; Zhaoyan Li; Siran Yang"
    props.keywords = "engineering monitoring; time-series forecasting; event management; provenance"
    props.comments = (
        "Generated from the GPT-reviewed scientific-consistency source for SoftwareX minor revision "
        "SOFTX-D-26-00931."
    )

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.widow_control = True

    style_settings = {
        "Title": (16, True, BLUE, 0, 8),
        "Heading 1": (13, True, BLACK, 10, 3),
        "Heading 2": (12, True, BLACK, 8, 2),
        "Heading 3": (11, True, BLACK, 6, 2),
        "Heading 4": (10.5, True, BLACK, 5, 1),
    }
    for name, (size, bold, color, before, after) in style_settings.items():
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Figure Caption" not in [style.name for style in document.styles]:
        caption = document.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = document.styles["Figure Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(9.5)
    caption.font.italic = False
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_with_next = False

    for name, left_indent, first_line in (
        ("List Number", Mm(7), Mm(-4)),
        ("List Bullet", Mm(7), Mm(-4)),
    ):
        if name not in [style.name for style in document.styles]:
            list_style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            list_style.base_style = normal
        else:
            list_style = document.styles[name]
        list_style.paragraph_format.left_indent = left_indent
        list_style.paragraph_format.first_line_indent = first_line
        list_style.paragraph_format.space_after = Pt(2)

    document.settings.odd_and_even_pages_header_footer = False
    for section in document.sections:
        section.different_first_page_header_footer = False
        section.page_width = Mm(215.9)
        section.page_height = Mm(279.4)
        section.top_margin = Mm(25.4)
        section.bottom_margin = Mm(25.4)
        section.left_margin = Mm(25.4)
        section.right_margin = Mm(25.4)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)
        sect_pr = section._sectPr
        line_number = sect_pr.find(qn("w:lnNumType"))
        if line_number is None:
            line_number = OxmlElement("w:lnNumType")
            sect_pr.append(line_number)
        line_number.set(qn("w:countBy"), "1")
        line_number.set(qn("w:restart"), "continuous")
        reset_footer(section.footer, include_page_number=True)
        reset_footer(section.first_page_footer, include_page_number=False)
        reset_footer(section.even_page_footer, include_page_number=False)


def configure_response_compaction(document: Document) -> None:
    """Keep the reviewer response readable while avoiding a near-empty trailing page."""
    normal = document.styles["Normal"]
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_after = Pt(2)

    for name, before, after in (
        ("Heading 1", 8, 2),
        ("Heading 2", 6, 1.5),
        ("Heading 3", 5, 1.5),
        ("Heading 4", 4, 1),
    ):
        style = document.styles[name]
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    relationship = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), MID_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.extend((color, underline))
    run.append(run_props)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(
    r"(<sup>.*?</sup>|\*\*.*?\*\*|`.*?`|<https?://[^>]+>|\*[^*]+?\*)"
)


def add_inline(paragraph, text: str, default_bold: bool = False, default_italic: bool = False) -> None:
    position = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            run.bold = default_bold
            run.italic = default_italic
        token = match.group(0)
        if token.startswith("<sup>"):
            run = paragraph.add_run(re.sub(r"</?sup>", "", token))
            run.font.superscript = True
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif token.startswith("<http"):
            add_hyperlink(paragraph, token[1:-1], token[1:-1])
        else:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        run.bold = default_bold
        run.italic = default_italic


def add_text_paragraph(document: Document, text: str, style: str | None = None, *, quote=False):
    paragraph = document.add_paragraph(style=style)
    if quote:
        paragraph.paragraph_format.left_indent = Mm(8)
        paragraph.paragraph_format.right_indent = Mm(4)
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(4)
        p_pr = paragraph._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), LIGHT_GRAY)
        p_pr.append(shading)
    add_inline(paragraph, text, default_italic=quote)
    return paragraph


def add_figure(document: Document, figure_number: int) -> None:
    image_path, caption, width = FIGURES[figure_number]
    if not image_path.exists():
        raise FileNotFoundError(f"Missing Figure {figure_number}: {image_path}")
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.paragraph_format.keep_together = True
    picture.paragraph_format.space_before = Pt(6)
    picture.paragraph_format.space_after = Pt(2)
    picture.add_run().add_picture(str(image_path), width=width)
    caption_paragraph = document.add_paragraph(style="Figure Caption")
    add_inline(caption_paragraph, caption)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows = []
    index = start
    while index < len(lines) and lines[index].lstrip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(rows):
        for col_index in range(columns):
            cell = table.cell(row_index, col_index)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            value = row[col_index] if col_index < len(row) else ""
            add_inline(paragraph, value, default_bold=(row_index == 0))
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(8.2 if columns >= 6 else 8.8)
    set_repeat_table_header(table.rows[0])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(document: Document, lines: Iterable[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F5F8")
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Courier New"
    run.font.size = Pt(7.7)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def normalized_heading(text: str) -> str:
    return re.sub(r"\s+", " ", re.sub(r"[*`_]", "", text)).strip()


def build_document(
    source: Path,
    output: Path,
    template: Path,
    document_kind: str,
    page_refs: dict[str, str] | None = None,
) -> None:
    document = Document(str(template))
    clear_body(document)
    title = (
        "SHM-EM: A forecast-aware event management framework for heterogeneous engineering monitoring"
        if document_kind == "manuscript"
        else "Response to Reviewers: SOFTX-D-26-00931"
    )
    configure_document(
        document,
        title,
        "SoftwareX revised manuscript" if document_kind == "manuscript" else "SoftwareX reviewer response",
    )
    if document_kind == "response":
        configure_response_compaction(document)

    lines = source.read_text(encoding="utf-8").splitlines()
    index = 0
    first_heading = True
    in_comment = False
    current_review_item: str | None = None
    append_location_to_next = False

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if stripped.startswith("<!--"):
            in_comment = True
        if in_comment:
            if stripped.endswith("-->"):
                in_comment = False
            index += 1
            continue
        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            code_lines = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            index += 1
            add_code_block(document, code_lines)
            continue

        figure_note = re.match(r">\s*\*\*Figure\s+(\d+) insertion note\.\*\*", stripped)
        if figure_note and document_kind == "manuscript":
            add_figure(document, int(figure_note.group(1)))
            index += 1
            continue

        heading = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = normalized_heading(heading.group(2))
            if first_heading:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(paragraph, text)
                first_heading = False
            else:
                paragraph = document.add_paragraph(style=f"Heading {min(level, 4)}")
                add_inline(paragraph, text)
            item_match = re.match(r"(R\d+-\d+)", text)
            if item_match:
                current_review_item = item_match.group(1)
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue

        if re.match(r"^\d+\.\s+", stripped):
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, stripped)
            index += 1
            continue
        if re.match(r"^[-*]\s+", stripped):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, "• " + re.sub(r"^[-*]\s+", "", stripped))
            index += 1
            continue

        if stripped.startswith(">"):
            add_text_paragraph(document, stripped.lstrip("> "), quote=True)
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                break
            if (
                candidate.startswith(("#", "|", "```", ">", "<!--"))
                or re.match(r"^\d+\.\s+", candidate)
                or re.match(r"^[-*]\s+", candidate)
            ):
                break
            paragraph_lines.append(candidate)
            index += 1
        text = " ".join(paragraph_lines)
        if text.startswith("<sup>a</sup>"):
            text = text.replace("China <sup>b</sup>", "China\n<sup>b</sup>", 1)

        if re.fullmatch(r"\*\*(Table|Listing|Algorithm)\s+.*\*\*", text):
            paragraph = document.add_paragraph(style="Figure Caption")
            add_inline(paragraph, text[2:-2])
            set_keep_with_next(paragraph)
        else:
            if document_kind == "response" and text == "**Changes in manuscript**":
                append_location_to_next = True
            paragraph = add_text_paragraph(document, text)
            if document_kind == "manuscript" and (
                paragraph.text.startswith("Ji'an Liao")
                or paragraph.text.startswith("a Key Laboratory")
                or paragraph.text.startswith("b Key Laboratory")
                or paragraph.text.startswith("* Corresponding author")
            ):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if document_kind == "manuscript" and paragraph.text.startswith("[http"):
                paragraph.paragraph_format.left_indent = Mm(5)
            if document_kind == "manuscript" and re.match(r"^\[\d+\]", paragraph.text):
                paragraph.paragraph_format.left_indent = Mm(6)
                paragraph.paragraph_format.first_line_indent = Mm(-6)
            if append_location_to_next and text != "**Changes in manuscript**":
                reference = (page_refs or {}).get(current_review_item or "")
                if reference:
                    run = paragraph.add_run(f" Final location: {reference}.")
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(BLUE)
                append_location_to_next = False

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))


def build_highlights(source: Path, output: Path, template: Path) -> None:
    document = Document(str(template))
    clear_body(document)
    document.core_properties.title = "Highlights"
    document.core_properties.subject = "SoftwareX revised manuscript highlights"
    document.core_properties.author = "Ji'an Liao; Zifa Wang; Dengke Zhao; Jianming Wang; Zhaoyan Li; Siran Yang"
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("Highlights")
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(16)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    for line in source.read_text(encoding="utf-8").splitlines():
        if not line.startswith("- "):
            continue
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Mm(6)
        paragraph.paragraph_format.first_line_indent = Mm(-4)
        paragraph.paragraph_format.space_after = Pt(6)
        run = paragraph.add_run("• " + line[2:])
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
    for section in document.sections:
        reset_footer(section.footer, include_page_number=False)
        reset_footer(section.first_page_footer, include_page_number=False)
        reset_footer(section.even_page_footer, include_page_number=False)
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--kind", choices=("manuscript", "response", "highlights"), required=True)
    parser.add_argument("--source", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--template", type=Path, default=DEFAULT_TEMPLATE)
    parser.add_argument("--page-refs", type=Path)
    args = parser.parse_args()
    if args.template is None:
        parser.error("Provide --template or set SHM_EM_SUBMITTED_TEMPLATE.")
    if args.kind == "highlights":
        build_highlights(args.source, args.output, args.template)
        print(args.output)
        return
    page_refs = None
    if args.page_refs:
        page_refs = json.loads(args.page_refs.read_text(encoding="utf-8"))
    build_document(args.source, args.output, args.template, args.kind, page_refs)
    print(args.output)


if __name__ == "__main__":
    main()
